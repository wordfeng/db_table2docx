from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Pt
import output


class Table2Doc:

    def __init__(self):
        self.doc = Document()
        self.table_style = self.doc.styles["Table Grid"]
        self.table_style.font.name = u'宋体'
        self.table_style._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.table_style.font.size = Pt(12)
        self.table = None

    def create_table(self, table_name: [str], cols_name: [list]):
        self.doc.add_heading(table_name, 2)
        self.table = self.doc.add_table(rows=1, cols=len(cols_name), style='Table Grid')
        hc = self.table.rows[0].cells
        for i in range(0, len(cols_name)):
            hc[i].text = cols_name[i]

    def insert(self, records: [tuple]):
        """table: [Table],
        records = (
            (1, "213", "123"),
            (2, "211213", "123")
        )
        """
        for record in records:
            row_cells = self.table.add_row().cells
            for i, cell in enumerate(record):
                row_cells[i].text = "NULL" if cell is None else str(cell)

    def save(self, file_name: [str]):
        self.doc.save(output.__file__.strip("__init__.py") + "%s.docx" % file_name)


if __name__ == '__main__':
    print(output.__file__.strip("__init__.py") + "%s.docx")
